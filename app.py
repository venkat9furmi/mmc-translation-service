import os
import io
import re
import time
import requests
from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Pt

app = Flask(__name__)

# ─── CONFIG ────────────────────────────────────────────────────────────────────
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_URL     = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL   = "openai/gpt-oss-120b"  # llama-3.3-70b-versatile deprecated June 2026

SKIP_EXACT = {
    "commentary", "time", "picture:sound", "picture: sound",
    "end", "[no commentary]", "time in", "clock into part 2",
    "break bumper out", "break bumper in", "still to come",
    "opening titles", "menu", "closing sequence",
    "semi-final, fixtures", "ucl final, fixture", "uwcl final, fixture",
    "no", "item", "dur", "r/t", "",
}

# ─── TRANSLATION-ONLY PROMPTS ─────────────────────────────────────────────────
# NOTE: All structural instructions (tables, cell maps, "keep English + add
# translation in same cell") were REMOVED on purpose. Python guarantees that
# mechanically — Groq never sees a table, so those instructions would be noise.
# Only LANGUAGE / STYLE / JUDGMENT rules remain — that is Groq's actual job.

GER_ROLE_PROMPT = """ROLE
You are a senior German football TV magazine script editor and voice-over writer (UEFA club competitions).

TASK
Translate and adapt the given English TV script lines into natural, professional German football broadcast language.

CONTEXT
This is a TV magazine / highlights script — NOT live match commentary.
The German must sound like it was written by a TV editor, not translated by AI.

WHAT TO TRANSLATE (you will only ever receive viewer-facing lines):
- Commentary / VO lines -> natural German broadcast narration
- SOT / UPSOT / NATSOT -> spoken, authentic German (not literal)
- On-screen text (only the visible string) -> translate only the visible text

LANGUAGE RULES (VERY IMPORTANT)
- Do NOT translate word-for-word
- Rewrite into natural German football TV language
- Use short, spoken sentences
- Use active voice
- Keep rhythm suitable for voice-over

STYLE (MMC / UEFA MAGAZINE)
- Controlled energy (not overdramatic)
- No cliches or PR language
- Clean, modern football phrasing
- Use terms like: K.o.-Phase, Achtelfinale, Hinspiel/Rueckspiel, Auswaerts, Treffer

QUALITY CONTROL
- Read every German line as spoken TV audio -- if it sounds translated, rewrite it
- If something is unclear, write [VERIFY]
- If the transcript itself is missing or just a placeholder, write [TRANSCRIPT MISSING]"""

FRA_ROLE_PROMPT = """ROLE
You are a senior French-language football TV magazine script editor and voice-over writer for UEFA club competitions, writing for French-speaking Switzerland.

TASK
Translate/adapt the given English TV script lines into natural, professional Swiss French football broadcast language.

CONTEXT
This is NOT live match commentary. The French must sound like polished Swiss French football TV writing, not AI translation.

SWISS FRENCH NUMBERS -- MANDATORY
- 70-79 = septante, septante et un, septante-deux, etc.
- 90-99 = nonante, nonante et un, nonante-deux, etc.
- Do NOT use "soixante-dix" or "quatre-vingt-dix"

WHAT TO TRANSLATE:
- VO / commentary lines
- SOT / UPSOT / NATSOT lines -- UNLESS the line's label is marked (FRENCH), see rule below
- fan vox pops, reporter questions, pundit lines
- visible GFX / on-screen text

(FRENCH)-LABELED ROWS -- CRITICAL RULE
You will be told the row's label (e.g. "SOT: VITINHA (FRENCH)") alongside the line.
- If the label contains "(FRENCH)", this is French-source audio.
- Do NOT back-translate it from the English transcript shown to you.
- Instead output exactly: FR: [TRANSCRIPT MISSING - FRENCH AUDIO]
- If the label does NOT contain "(FRENCH)", translate normally as usual.

PRONUNCIATION NOTES
Add a pronunciation note in brackets only for difficult foreign player/club names, on first useful mention.
Example: Joao Neves [prononciation : Jo-ao Ne-vech]
Do NOT add pronunciation notes for obvious names (Arsenal, Paris, Lyon, Inter, Real Madrid, Harry Kane, etc).

STYLE
- Spoken, clear, concise; active voice; controlled energy; modern broadcast tone
- No word-for-word translation, no English sentence structure
- No PR language, no newspaper style, no overdramatic cliches
- Natural football terms: match aller / match retour, a domicile / a l'exterieur, but,
  egalisation, qualification, finale, demi-finale, temps additionnel, arrets de jeu,
  seance de tirs au but

CLUB NAMES
- Bayern Muenchen -> Bayern Munich | Arsenal FC -> Arsenal
- Barcelona -> Barcelone or le Barca (avoid "La Blaugrana"; prefer "le Barca"/"les Blaugrana")
- Paris Saint-Germain -> Paris Saint-Germain or Paris
- Inter Milan -> l'Inter Milan or l'Inter
- Lyon / OL Lyonnes -> Lyon or l'OL (unless official wording requires otherwise)

SPECIFIC FIXES
- "Ball Knowledge" -> "connaissances en football" / "connaissances football" / "culture foot"
- "I would give it a seven" (self-rating) -> "Je me donnerais un sept."
- Do NOT use "tout etait a refaire" for "back to the drawing board" unless the team
  actually lost an advantage
- Avoid "au final"; use "finalement", "au bout du compte", "en realite", or restructure
- Avoid repeating "finale/final" too close together
- For "bridesmaids", do NOT use "juste derriere"; use "passes tout pres" or
  "echoue a la derniere marche"
- Use "la pression de gagner", not "la pression pour gagner"

QUALITY CONTROL
- Read every French line as spoken TV audio -- if it sounds translated, rewrite it
- If unclear, write FR: [VERIFY]
- If the (non-French) transcript itself is missing, write FR: [TRANSCRIPT MISSING]"""


# ─── HELPERS ───────────────────────────────────────────────────────────────────

def cell_text(cell):
    return "\n".join(p.text for p in cell.paragraphs if p.text.strip())


def should_skip(text):
    t = text.strip()
    if not t:
        return True
    if t.lower() in SKIP_EXACT:
        return True
    if re.match(r"^\d{2}:\d{2}:\d{2}$", t):
        return True
    return False


def build_cell_map(doc):
    """
    Walk every table/row and build a CELL MAP of translatable cells.
    Key format: T{table}R{row}C{col}

    Each value is a dict: {"text": ..., "label": ...}
    "label" is the row's column-1 tag (e.g. "SOT: VITINHA (FRENCH)") so the
    model can apply the (FRENCH)-source rule without ever touching structure.
    """
    cell_map = {}
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            n_cols = len(row.cells)

            # Running order table (Table 0): translate FEATURE descriptions in col 1
            if t_idx == 0 and n_cols >= 2:
                text = cell_text(row.cells[1])
                if not should_skip(text) and "FEATURE" in text.upper():
                    cell_map[f"T{t_idx}R{r_idx}C1"] = {"text": text, "label": ""}

            # Feature tables: translate COMMENTARY col (usually col 3)
            # label = column 1 of the SAME row (PICTURE:SOUND / SOT tag)
            if t_idx >= 1 and n_cols >= 4:
                text  = cell_text(row.cells[3])
                label = cell_text(row.cells[1]) if n_cols > 1 else ""
                if not should_skip(text):
                    cell_map[f"T{t_idx}R{r_idx}C3"] = {"text": text, "label": label}

    return cell_map


def build_single_prompt(cell_map, language):
    role_prompt = GER_ROLE_PROMPT if language == "de" else FRA_ROLE_PROMPT
    prefix      = "DE:" if language == "de" else "FR:"

    lines = "\n".join(
        f"{key} [label: {v['label']}]: {v['text']}" if v["label"] else f"{key}: {v['text']}"
        for key, v in cell_map.items()
    )

    return f"""{role_prompt}

INPUT FORMAT NOTE
Each input line starts with a KEY (e.g. T2R4C3). Some lines also show
"[label: ...]" — that label is ONLY for you to detect the (FRENCH)-source
rule. It is NOT part of the text to translate and must NEVER appear in your output.

OUTPUT FORMAT — MANDATORY:
For every line below, output it EXACTLY like this (KEY only, no brackets, no label):
KEY: {prefix} <translation>

Do NOT repeat the English text. Do NOT add explanations, headers, or summaries.
Output ONLY the KEY: {prefix} <translation> lines — one per input line, in the same order.
If a line cannot be reliably translated, output: KEY: {prefix} [VERIFY]

INPUT (key [optional label]: original English line):
{lines}"""


def call_groq(prompt, api_key, max_retries=2, max_wait=8):
    """
    Calls Groq with sane token sizing + bounded 429 backoff.

    Free tier for openai/gpt-oss-120b = 8,000 TOKENS PER MINUTE (not per call).
    A batch of 20 cells needs ~600-1200 output tokens, never 8000 — so we
    right-size the request instead of grabbing the whole per-minute budget
    on a single call.

    IMPORTANT: we cap how long we'll sleep on a 429. Groq's Retry-After can
    sometimes be 30-60s+ — blindly sleeping that long INSIDE an HTTP request
    causes Railway/n8n's own request timeout to kill the connection first,
    producing a generic content-less 500 (worse than just failing fast).
    Better to give up quickly and let the backfill/[VERIFY] safety net in
    translate_full_script handle it visibly.
    """
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": GROQ_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
        "max_tokens": 1536,  # right-sized for a 20-cell batch, not 8000
    }

    print(f"[DEBUG] Groq request: prompt={len(prompt)} chars, "
          f"api_key_len={len(api_key)}, api_key_preview={api_key[:8]}...",
          flush=True)

    for attempt in range(max_retries):
        resp = requests.post(GROQ_URL, headers=headers, json=payload, timeout=60)
        if resp.status_code == 429:
            retry_after = min(float(resp.headers.get("Retry-After", 3)), max_wait)
            if attempt < max_retries - 1:
                time.sleep(retry_after)
                continue
            return ""  # give up fast — caller treats missing keys as [VERIFY]
        if resp.status_code >= 400:
            print(f"[DEBUG] Groq error {resp.status_code}: {resp.text[:500]}", flush=True)
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]

    return ""


def parse_table_map_response(output, keys):
    """
    Parse Groq's response in the format:
        T2R4C3: DE: Translated text here
    Back into {key: translated_text}
    """
    translations = {}
    key_set = set(keys)

    for line in output.splitlines():
        line = line.strip()
        if not line or ":" not in line:
            continue
        # Split off the key (everything before first colon)
        possible_key, _, rest = line.partition(":")
        possible_key = possible_key.strip()
        if possible_key in key_set:
            translations[possible_key] = rest.strip()

    return translations


BATCH_SIZE = 20          # cells per Groq call — small enough to stay reliable
MAX_BACKFILL_ROUNDS = 2  # retry rounds for any keys the model dropped


def translate_full_script(cell_map, language, api_key):
    """
    Batched translation with a completeness guarantee.

    Why not one giant call? Long structured lists (50+ numbered items)
    are a known LLM failure mode — models silently skip items or drift
    on key formatting, with no error thrown. So we:

      1. Send small batches (BATCH_SIZE cells) — keeps quality + key
         formatting reliable per call.
      2. After all batches, diff what we asked for vs what we got back.
      3. Anything missing gets a focused backfill call (repeated up to
         MAX_BACKFILL_ROUNDS times).
      4. Anything STILL missing after that is marked [VERIFY] so it is
         visible in the human review step — never silently dropped.
    """
    all_keys      = list(cell_map.keys())
    translations  = {}

    # ── Pass 1: batched translation ──────────────────────────────────────
    for i, start in enumerate(range(0, len(all_keys), BATCH_SIZE)):
        if i > 0:
            time.sleep(2)  # spread calls across the per-minute token window
        batch_keys = all_keys[start : start + BATCH_SIZE]
        batch      = {k: cell_map[k] for k in batch_keys}
        prompt     = build_single_prompt(batch, language)
        output     = call_groq(prompt, api_key)
        translations.update(parse_table_map_response(output, batch_keys))

    # ── Pass 2: backfill any keys the model silently dropped ─────────────
    for _ in range(MAX_BACKFILL_ROUNDS):
        missing = [k for k in all_keys if k not in translations]
        if not missing:
            break
        time.sleep(2)
        batch  = {k: cell_map[k] for k in missing}
        prompt = build_single_prompt(batch, language)
        output = call_groq(prompt, api_key)
        translations.update(parse_table_map_response(output, missing))

    # ── Pass 3: anything still missing → flag for human review ───────────
    prefix = "DE:" if language == "de" else "FR:"
    for k in all_keys:
        if k not in translations:
            translations[k] = "[VERIFY]"

    return translations


def add_translation_to_cell(cell, translation_text, prefix):
    if not translation_text:
        return
    para = cell.add_paragraph()
    run  = para.add_run(f"{prefix} {translation_text}")
    run.font.size = Pt(10)
    run.italic = True


# ─── ROUTES ────────────────────────────────────────────────────────────────────

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "service": "MMC Sport Translation Service v2"})


@app.route("/translate", methods=["POST"])
def translate():
    try:
        api_key  = request.form.get("api_key") or GROQ_API_KEY
        language = request.form.get("language", "de").lower()

        if not api_key:
            return jsonify({"error": "GROQ_API_KEY missing"}), 400
        if language not in ("de", "fr"):
            return jsonify({"error": "language must be 'de' or 'fr'"}), 400
        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if not file.filename.lower().endswith(".docx"):
            return jsonify({"error": "Only .docx files are supported"}), 400

        doc_bytes = file.read()
        doc       = Document(io.BytesIO(doc_bytes))

        # ── Step 1: Build the cell map ───────────────────────────────────────
        cell_map = build_cell_map(doc)
        if not cell_map:
            return jsonify({"error": "No translatable content found"}), 400

        # ── Step 2: Single (or minimal) call to Groq ─────────────────────────
        translations = translate_full_script(cell_map, language, api_key)

        # ── Step 3: Apply translations to a fresh doc copy ───────────────────
        doc    = Document(io.BytesIO(doc_bytes))
        prefix = "DE:" if language == "de" else "FR:"

        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                n_cols = len(row.cells)

                if t_idx == 0 and n_cols >= 2:
                    key = f"T{t_idx}R{r_idx}C1"
                    if key in translations:
                        add_translation_to_cell(row.cells[1], translations[key], prefix)

                if t_idx >= 1 and n_cols >= 4:
                    key = f"T{t_idx}R{r_idx}C3"
                    if key in translations:
                        add_translation_to_cell(row.cells[3], translations[key], prefix)

        # ── Step 4: Save & return ─────────────────────────────────────────────
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        lang_tag    = "GER" if language == "de" else "FRA"
        base_name   = file.filename.replace(".docx", "")
        output_name = f"{base_name}_{lang_tag}_translated.docx"

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=output_name,
        )

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Groq API call failed: {str(e)}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
